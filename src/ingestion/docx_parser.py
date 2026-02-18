"""Word document parser for contract extraction.

Extracts text, tables, and structural information from .docx files
using python-docx. Handles both file paths and raw bytes.
"""

import io
from pathlib import Path

from docx import Document as DocxDocument


def _open_docx(file_path_or_bytes) -> DocxDocument:
    """Open a docx file from path or bytes."""
    if isinstance(file_path_or_bytes, (str, Path)):
        return DocxDocument(str(file_path_or_bytes))
    elif isinstance(file_path_or_bytes, (bytes, bytearray)):
        return DocxDocument(io.BytesIO(file_path_or_bytes))
    else:
        raise ValueError(
            f"Expected file path (str/Path) or bytes, got {type(file_path_or_bytes)}"
        )


def extract_raw_text(file_path_or_bytes) -> str:
    """Extract all text from a Word document, preserving paragraph breaks.

    Args:
        file_path_or_bytes: Path to .docx file or raw bytes.

    Returns:
        Full text content with paragraphs separated by newlines.
    """
    doc = _open_docx(file_path_or_bytes)
    paragraphs = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            paragraphs.append(text)
    return "\n".join(paragraphs)


def extract_tables(file_path_or_bytes) -> list[dict]:
    """Extract all tables from a Word document.

    Each table is returned as a dict with:
        - headers: list of column header strings
        - rows: list of row dicts mapping header -> cell value
        - raw_rows: list of lists (raw cell text per row)

    Args:
        file_path_or_bytes: Path to .docx file or raw bytes.

    Returns:
        List of table dicts.
    """
    doc = _open_docx(file_path_or_bytes)
    tables = []

    for table in doc.tables:
        raw_rows = []
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            raw_rows.append(cells)

        if not raw_rows:
            continue

        # First row is typically headers
        headers = raw_rows[0]
        rows = []
        for raw_row in raw_rows[1:]:
            row_dict = {}
            for i, cell in enumerate(raw_row):
                key = headers[i] if i < len(headers) else f"col_{i}"
                row_dict[key] = cell
            rows.append(row_dict)

        tables.append({
            "headers": headers,
            "rows": rows,
            "raw_rows": raw_rows,
        })

    return tables


def extract_headings(file_path_or_bytes) -> list[dict]:
    """Extract heading paragraphs with their level and text.

    Identifies headings by Word heading styles (Heading 1, Heading 2, etc.)
    and also detects bold-only paragraphs as implicit headings.

    Returns:
        List of dicts with 'level', 'text', and 'style' keys.
    """
    doc = _open_docx(file_path_or_bytes)
    headings = []

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue

        style_name = para.style.name if para.style else ""

        # Check for explicit heading styles
        if style_name.startswith("Heading"):
            try:
                level = int(style_name.replace("Heading ", "").strip())
            except ValueError:
                level = 1
            headings.append({"level": level, "text": text, "style": style_name})

        # Check for bold-only paragraphs (implicit headings)
        elif len(text) < 100 and para.runs:
            all_bold = all(run.bold for run in para.runs if run.text.strip())
            if all_bold and para.runs:
                headings.append({"level": 3, "text": text, "style": "bold_implicit"})

    return headings


def extract_full_content(file_path_or_bytes) -> dict:
    """Extract complete document content: text, tables, and structure.

    This is the main entry point for the hybrid extractor. Returns
    all content needed for both regex and LLM extraction.

    Args:
        file_path_or_bytes: Path to .docx file or raw bytes.

    Returns:
        Dict with:
            - full_text: All text content joined with newlines
            - tables: List of extracted table dicts
            - headings: List of heading dicts
            - table_texts: List of table content as formatted text strings
    """
    full_text = extract_raw_text(file_path_or_bytes)
    tables = extract_tables(file_path_or_bytes)
    headings = extract_headings(file_path_or_bytes)

    # Convert tables to readable text for inclusion in extraction
    table_texts = []
    for table in tables:
        lines = []
        if table["headers"]:
            lines.append(" | ".join(table["headers"]))
            lines.append("-" * len(lines[0]))
        for row in table["raw_rows"][1:]:
            lines.append(" | ".join(row))
        table_texts.append("\n".join(lines))

    # Combine text + table content for comprehensive extraction
    combined_text = full_text
    if table_texts:
        combined_text += "\n\n--- TABLES ---\n\n" + "\n\n".join(table_texts)

    return {
        "full_text": full_text,
        "combined_text": combined_text,
        "tables": tables,
        "headings": headings,
        "table_texts": table_texts,
    }


def validate_docx(file_path_or_bytes) -> tuple[bool, str]:
    """Validate that a file is a readable Word document with content.

    Returns:
        Tuple of (is_valid, error_message).
    """
    try:
        doc = _open_docx(file_path_or_bytes)
    except Exception as e:
        return False, f"Cannot read file as a Word document: {e}"

    # Check that it has some text content
    text_content = ""
    for para in doc.paragraphs:
        text_content += para.text
        if len(text_content) > 100:
            break

    if len(text_content.strip()) < 50:
        return False, (
            "Word document appears to be empty or contains very little text. "
            "Please upload a contract document with substantive content."
        )

    return True, ""
