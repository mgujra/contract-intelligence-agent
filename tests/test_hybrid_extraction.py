"""Tests for the hybrid extraction pipeline (docx parser, regex, LLM, orchestrator).

Creates synthetic .docx files in test setup to validate the full pipeline
from Word document upload to vector-store-ready text.
"""

import io
import json
import tempfile
from pathlib import Path
from unittest.mock import MagicMock, patch

import pytest
from docx import Document as DocxDocument

# ── Helpers ──────────────────────────────────────────────────────────────────

def _create_test_docx(paragraphs: list[str] | None = None, tables: list[list[list[str]]] | None = None) -> bytes:
    """Create a .docx file in memory and return raw bytes.

    Args:
        paragraphs: List of paragraph strings to add.
        tables: List of tables, each table is a list of rows, each row is a list of cells.
    """
    doc = DocxDocument()

    if paragraphs:
        for text in paragraphs:
            doc.add_paragraph(text)

    if tables:
        for table_data in tables:
            if not table_data:
                continue
            num_cols = len(table_data[0])
            table = doc.add_table(rows=len(table_data), cols=num_cols)
            for i, row_data in enumerate(table_data):
                for j, cell_text in enumerate(row_data):
                    table.rows[i].cells[j].text = cell_text

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


SAMPLE_CONTRACT_TEXT = """
CONTRACT NUMBER: W911NF-25-D-0042

CONTRACT TYPE: Firm-Fixed-Price (FFP)
CONTRACTING AGENCY: Department of the Army (ARMY)
CONTRACTING OFFICER: Col. Jane Smith

CONTRACTOR INFORMATION
Contractor: Apex Defense Technologies LLC
CAGE Code: 5AB12
DUNS/UEI: 987654321ABC
Business Size: Small

NAICS CODE: 541512
PSC CODE: D318

TOTAL CONTRACT VALUE
Base Period Value: $18,172,478.80
Contract Ceiling: $45,000,000.00
Funded Amount: $12,500,000.00

PERIOD OF PERFORMANCE
Base Period: 2025-01-15 through 2026-01-14 (12 months)
Option Period 1: 2026-01-15 through 2027-01-14 (12 months)

STATEMENT OF WORK
The contractor shall provide cybersecurity monitoring and incident response services
for Army network infrastructure. Services include 24/7 SOC operations, threat hunting,
vulnerability assessments, and NIST SP 800-171 compliance support.

CONTRACT LINE ITEMS (CLINs)
CLIN 0001: Cybersecurity Monitoring Services
Type: FFP | Qty: 1 Lot | Unit Price: $10,000,000.00 | Total: $10,000,000.00
CLIN 0002: Incident Response Retainer
Type: FFP | Qty: 12 Month | Unit Price: $681,039.90 | Total: $8,172,478.80

APPLICABLE FAR CLAUSES
FAR 52.204-21 - Basic Safeguarding of Covered Contractor Information Systems
FAR 52.232-40 - Providing Accelerated Payments to Small Business Subcontractors

APPLICABLE DFARS CLAUSES
DFARS 252.204-7012 - Safeguarding Covered Defense Information [MANDATORY FLOWDOWN]
DFARS 252.204-7020 - NIST SP 800-171 DoD Assessment Requirements

SECURITY REQUIREMENTS
- SECRET clearance required for all key personnel
- CMMC Level 2 certification required
- NIST SP 800-171 compliance required

SPECIAL CONTRACT PROVISIONS
- Key Personnel: Program Manager, SOC Lead, Incident Response Lead
- Monthly status reporting required
- Government-furnished equipment provided (GFE)
- DevSecOps pipeline implementation required

SUBCONTRACTOR REQUIREMENTS
Subcontractor: SecureNet Analytics (CAGE: 7CD34)
Business Size: Small
Estimated Value: $2,500,000.00
Scope: Threat intelligence and analytics platform
Flowdown Clauses: DFARS 252.204-7012, DFARS 252.204-7020
"""


def _create_contract_docx() -> bytes:
    """Create a realistic contract .docx for integration tests."""
    return _create_test_docx(paragraphs=SAMPLE_CONTRACT_TEXT.strip().split("\n"))


# ══════════════════════════════════════════════════════════════════════════════
# DOCX PARSER TESTS
# ══════════════════════════════════════════════════════════════════════════════

class TestDocxParser:
    """Tests for src.ingestion.docx_parser."""

    def test_extract_raw_text(self):
        from src.ingestion.docx_parser import extract_raw_text

        docx_bytes = _create_test_docx(paragraphs=["Hello World", "Second paragraph"])
        text = extract_raw_text(docx_bytes)
        assert "Hello World" in text
        assert "Second paragraph" in text

    def test_extract_tables(self):
        from src.ingestion.docx_parser import extract_tables

        table = [
            ["Header A", "Header B"],
            ["Row1-A", "Row1-B"],
            ["Row2-A", "Row2-B"],
        ]
        docx_bytes = _create_test_docx(tables=[table])
        tables = extract_tables(docx_bytes)
        assert len(tables) == 1
        assert tables[0]["headers"] == ["Header A", "Header B"]
        assert len(tables[0]["rows"]) == 2
        assert tables[0]["rows"][0]["Header A"] == "Row1-A"

    def test_extract_full_content(self):
        from src.ingestion.docx_parser import extract_full_content

        docx_bytes = _create_test_docx(
            paragraphs=["Intro text"],
            tables=[[["Col1", "Col2"], ["Data1", "Data2"]]],
        )
        content = extract_full_content(docx_bytes)
        assert "Intro text" in content["full_text"]
        assert "combined_text" in content
        assert len(content["tables"]) == 1
        # Combined text should include both paragraph and table text
        assert "Intro text" in content["combined_text"]
        assert "Data1" in content["combined_text"]

    def test_validate_docx_valid(self):
        from src.ingestion.docx_parser import validate_docx

        docx_bytes = _create_test_docx(paragraphs=[
            "This is a valid contract document with enough substantive content to pass validation.",
            "It contains multiple paragraphs and exceeds the minimum character threshold required.",
        ])
        is_valid, error = validate_docx(docx_bytes)
        assert is_valid is True
        assert error == ""

    def test_validate_docx_invalid_bytes(self):
        from src.ingestion.docx_parser import validate_docx

        is_valid, error = validate_docx(b"not a docx file")
        assert is_valid is False
        assert "valid Word" in error or "corrupt" in error.lower() or "not a valid" in error.lower()

    def test_validate_docx_empty(self):
        from src.ingestion.docx_parser import validate_docx

        # Create a valid but empty docx
        doc = DocxDocument()
        buf = io.BytesIO()
        doc.save(buf)
        is_valid, error = validate_docx(buf.getvalue())
        assert is_valid is False
        assert "empty" in error.lower() or "no text" in error.lower()

    def test_extract_headings(self):
        from src.ingestion.docx_parser import extract_headings

        doc = DocxDocument()
        doc.add_heading("Main Title", level=1)
        doc.add_paragraph("Some body text")
        doc.add_heading("Subsection", level=2)
        buf = io.BytesIO()
        doc.save(buf)

        headings = extract_headings(buf.getvalue())
        assert len(headings) >= 2
        assert any(h["text"] == "Main Title" for h in headings)
        assert any(h["text"] == "Subsection" for h in headings)

    def test_extract_from_file_path(self):
        from src.ingestion.docx_parser import extract_raw_text

        docx_bytes = _create_test_docx(paragraphs=["File path test"])
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
            f.write(docx_bytes)
            f.flush()
            text = extract_raw_text(f.name)
        assert "File path test" in text


# ══════════════════════════════════════════════════════════════════════════════
# REGEX EXTRACTOR TESTS
# ══════════════════════════════════════════════════════════════════════════════

class TestRegexExtractors:
    """Tests for src.ingestion.regex_extractors."""

    def test_extract_contract_number_labeled(self):
        from src.ingestion.regex_extractors import extract_contract_number

        result = extract_contract_number("CONTRACT NUMBER: W911NF-25-D-0042\nMore text")
        assert result.value == "W911NF-25-D-0042"
        assert result.confidence >= 0.9

    def test_extract_contract_number_federal_format(self):
        from src.ingestion.regex_extractors import extract_contract_number

        result = extract_contract_number("Agreement ref: FA8721-23-C-1234 dated Jan 2025")
        assert result.value == "FA8721-23-C-1234"

    def test_extract_dollar_values(self):
        from src.ingestion.regex_extractors import extract_dollar_values

        text = "Base Period Value: $18,172,478.80\nContract Ceiling: $45,000,000.00\nFunded Amount: $12,500,000.00"
        values = extract_dollar_values(text)
        assert "value" in values
        assert values["value"].value == pytest.approx(18172478.80, rel=0.01)
        assert "ceiling_value" in values
        assert values["ceiling_value"].value == pytest.approx(45000000.00, rel=0.01)
        assert "funded_amount" in values
        assert values["funded_amount"].value == pytest.approx(12500000.00, rel=0.01)

    def test_extract_dollar_values_shorthand(self):
        from src.ingestion.regex_extractors import extract_dollar_values

        text = "Total value: $1.5M with ceiling of $5M"
        values = extract_dollar_values(text)
        assert any(v.value == pytest.approx(1500000, rel=0.01) for v in values.values())

    def test_extract_dates(self):
        from src.ingestion.regex_extractors import extract_dates

        text = "Base Period: 2025-01-15 through 2026-01-14 (12 months)"
        dates = extract_dates(text)
        assert "base_period_start" in dates
        assert "2025-01-15" in dates["base_period_start"].value

    def test_extract_far_clauses(self):
        from src.ingestion.regex_extractors import extract_far_clauses

        text = """
        FAR 52.204-21 - Basic Safeguarding of Covered Contractor Information Systems
        FAR 52.232-40 - Providing Accelerated Payments to Small Business Subcontractors
        """
        result = extract_far_clauses(text)
        assert result.value is not None
        assert len(result.value) == 2
        assert result.value[0]["number"] == "52.204-21"

    def test_extract_dfars_clauses(self):
        from src.ingestion.regex_extractors import extract_dfars_clauses

        text = """
        DFARS 252.204-7012 - Safeguarding Covered Defense Information [MANDATORY FLOWDOWN]
        DFARS 252.204-7020 - NIST SP 800-171 DoD Assessment Requirements
        """
        result = extract_dfars_clauses(text)
        assert result.value is not None
        assert len(result.value) == 2
        assert result.value[0]["number"] == "252.204-7012"
        assert result.value[0].get("flowdown") is True

    def test_extract_contractor_info(self):
        from src.ingestion.regex_extractors import extract_contractor_info

        text = "Contractor: Apex Defense Technologies LLC\nCAGE Code: 5AB12\nBusiness Size: Small"
        info = extract_contractor_info(text)
        assert info["contractor_name"].value == "Apex Defense Technologies LLC"
        assert info["cage"].value == "5AB12"
        assert info["size"].value == "Small"

    def test_extract_agency(self):
        from src.ingestion.regex_extractors import extract_agency

        text = "CONTRACTING AGENCY: Department of the Army (ARMY)"
        result = extract_agency(text)
        assert result.value is not None
        assert result.value["code"] == "ARMY"

    def test_extract_contract_type(self):
        from src.ingestion.regex_extractors import extract_contract_type

        text = "CONTRACT TYPE: Firm-Fixed-Price (FFP)"
        result = extract_contract_type(text)
        assert result.value == "FFP"

    def test_extract_contract_type_idiq(self):
        from src.ingestion.regex_extractors import extract_contract_type

        text = "This is an Indefinite Delivery/Indefinite Quantity contract"
        result = extract_contract_type(text)
        assert result.value == "IDIQ"

    def test_extract_naics_code(self):
        from src.ingestion.regex_extractors import extract_naics_code

        text = "NAICS CODE: 541512"
        result = extract_naics_code(text)
        assert result.value == "541512"

    def test_extract_psc_code(self):
        from src.ingestion.regex_extractors import extract_psc_code

        text = "PSC CODE: D318"
        result = extract_psc_code(text)
        assert result.value == "D318"

    def test_extract_all_regex(self):
        from src.ingestion.regex_extractors import extract_all_regex

        results = extract_all_regex(SAMPLE_CONTRACT_TEXT)
        assert "contract_number" in results
        assert results["contract_number"].value == "W911NF-25-D-0042"
        assert "contract_type" in results
        assert "value" in results

    def test_extract_contracting_officer(self):
        from src.ingestion.regex_extractors import extract_contracting_officer

        text = "CONTRACTING OFFICER: Col. Jane Smith"
        result = extract_contracting_officer(text)
        assert result.value == "Col. Jane Smith"

    def test_extract_subcontractors(self):
        from src.ingestion.regex_extractors import extract_subcontractors

        text = """
        Subcontractor: SecureNet Analytics (CAGE: 7CD34)
        Business Size: Small
        Estimated Value: $2,500,000.00
        Scope: Threat intelligence and analytics platform
        Flowdown Clauses: DFARS 252.204-7012, DFARS 252.204-7020
        """
        result = extract_subcontractors(text)
        assert result.value is not None
        assert len(result.value) >= 1
        assert result.value[0]["subcontractor_name"] == "SecureNet Analytics"


# ══════════════════════════════════════════════════════════════════════════════
# LLM EXTRACTOR TESTS
# ══════════════════════════════════════════════════════════════════════════════

class TestLLMExtractors:
    """Tests for src.ingestion.llm_extractors."""

    def test_mock_mode_scope(self):
        from src.ingestion.llm_extractors import extract_freeform_fields

        result = extract_freeform_fields(SAMPLE_CONTRACT_TEXT, mode="mock")
        assert result["confidence"] == 0.3  # Mock mode low confidence
        # Should detect scope-related keywords
        assert result["scope_of_work"] is not None or result["scope_of_work"] == ""

    def test_mock_mode_security(self):
        from src.ingestion.llm_extractors import extract_freeform_fields

        text = "SECRET clearance required. CMMC Level 2 certification required. NIST SP 800-171 compliance."
        result = extract_freeform_fields(text, mode="mock")
        assert len(result["security_requirements"]) >= 1

    def test_mock_mode_provisions(self):
        from src.ingestion.llm_extractors import extract_freeform_fields

        text = "Key Personnel: PM, Technical Lead. Monthly status report required. DevSecOps implementation."
        result = extract_freeform_fields(text, mode="mock")
        assert len(result["special_provisions"]) >= 1

    def test_mock_mode_domain_cybersecurity(self):
        from src.ingestion.llm_extractors import extract_freeform_fields

        text = "cybersecurity vulnerability penetration test incident response SOC SIEM zero trust"
        result = extract_freeform_fields(text, mode="mock")
        assert result["domain"] == "cybersecurity"

    def test_mock_mode_domain_it_modernization(self):
        from src.ingestion.llm_extractors import extract_freeform_fields

        text = "cloud migration digital transformation microservices agile containerization kubernetes"
        result = extract_freeform_fields(text, mode="mock")
        assert result["domain"] == "it_modernization"

    def test_mock_mode_returns_all_fields(self):
        from src.ingestion.llm_extractors import extract_freeform_fields

        result = extract_freeform_fields("some generic text", mode="mock")
        assert "scope_of_work" in result
        assert "security_requirements" in result
        assert "special_provisions" in result
        assert "domain" in result
        assert "confidence" in result

    def test_anthropic_mode_without_key_falls_back(self):
        from src.ingestion.llm_extractors import extract_freeform_fields

        # Without API key, should fall back to mock mode
        result = extract_freeform_fields("test text", mode="anthropic", api_key=None)
        assert result["confidence"] == 0.3  # Falls back to keyword mode


# ══════════════════════════════════════════════════════════════════════════════
# HYBRID EXTRACTOR TESTS
# ══════════════════════════════════════════════════════════════════════════════

class TestHybridExtractor:
    """Tests for src.ingestion.hybrid_extractor."""

    def test_extract_from_docx_basic(self):
        from src.ingestion.hybrid_extractor import extract_contract_from_docx

        docx_bytes = _create_contract_docx()
        result = extract_contract_from_docx(docx_bytes, "test_contract.docx", mode="mock")

        assert result.is_valid is True
        assert result.contract_data["contract_number"] == "W911NF-25-D-0042"
        assert result.contract_data["contract_type"] == "FFP"
        assert result.document_text  # Should have generated text
        assert "CONTRACT NUMBER:" in result.document_text

    def test_extract_report_populated(self):
        from src.ingestion.hybrid_extractor import extract_contract_from_docx

        docx_bytes = _create_contract_docx()
        result = extract_contract_from_docx(docx_bytes, "test.docx", mode="mock")

        report = result.report
        assert report.total_fields > 0
        assert report.regex_extracted > 0
        assert len(report.field_sources) > 0

    def test_extract_invalid_docx(self):
        from src.ingestion.hybrid_extractor import extract_contract_from_docx

        result = extract_contract_from_docx(b"invalid data", "bad.docx", mode="mock")
        assert result.is_valid is False
        assert result.error

    def test_merge_strategy_regex_wins_for_structured(self):
        from src.ingestion.hybrid_extractor import extract_contract_from_docx

        docx_bytes = _create_contract_docx()
        result = extract_contract_from_docx(docx_bytes, "test.docx", mode="mock")

        # Structured fields should come from regex
        sources = result.report.field_sources
        assert sources.get("contract_number") == "regex"
        assert sources.get("contract_type") == "regex"

    def test_merge_strategy_llm_wins_for_freeform(self):
        from src.ingestion.hybrid_extractor import extract_contract_from_docx

        docx_bytes = _create_contract_docx()
        result = extract_contract_from_docx(docx_bytes, "test.docx", mode="mock")

        # Check that freeform fields come from LLM
        sources = result.report.field_sources
        # If LLM found anything, it should be tagged as "llm"
        if result.contract_data.get("security_requirements"):
            assert sources.get("security_requirements") == "llm"

    def test_document_text_has_section_headers(self):
        from src.ingestion.hybrid_extractor import extract_contract_from_docx

        docx_bytes = _create_contract_docx()
        result = extract_contract_from_docx(docx_bytes, "test.docx", mode="mock")

        text = result.document_text
        assert "--- CONTRACTOR INFORMATION ---" in text
        assert "--- TOTAL CONTRACT VALUE ---" in text
        assert "--- STATEMENT OF WORK ---" in text
        assert "--- END OF CONTRACT ---" in text

    def test_document_text_compatible_with_chunker(self):
        """Verify that generated document_text can be chunked."""
        from src.ingestion.hybrid_extractor import extract_contract_from_docx
        from src.ingestion.chunker import chunk_contract

        docx_bytes = _create_contract_docx()
        result = extract_contract_from_docx(docx_bytes, "test.docx", mode="mock")

        chunks = chunk_contract(result.document_text)
        assert len(chunks) > 0
        # Chunks should have the contract number
        assert any(c.contract_number == "W911NF-25-D-0042" for c in chunks)

    def test_save_extraction_results(self):
        from src.ingestion.hybrid_extractor import extract_contract_from_docx, save_extraction_results

        docx_bytes = _create_contract_docx()
        result = extract_contract_from_docx(docx_bytes, "test.docx", mode="mock")

        with tempfile.TemporaryDirectory() as tmpdir:
            json_dir = Path(tmpdir) / "json"
            text_dir = Path(tmpdir) / "text"

            paths = save_extraction_results(result, json_directory=json_dir, text_directory=text_dir)

            assert paths["json_path"].exists()
            assert paths["text_path"].exists()

            # Verify JSON is valid
            json_data = json.loads(paths["json_path"].read_text())
            assert json_data["contract_number"] == "W911NF-25-D-0042"
            assert "document_text" in json_data

            # Verify text matches
            text_content = paths["text_path"].read_text()
            assert "CONTRACT NUMBER: W911NF-25-D-0042" in text_content

    def test_defaults_filled_for_missing_fields(self):
        """Test that missing fields get sensible defaults."""
        from src.ingestion.hybrid_extractor import extract_contract_from_docx

        # Minimal docx with no recognizable fields
        docx_bytes = _create_test_docx(paragraphs=[
            "This is a generic document without standard contract fields.",
            "It has some content but no structured data.",
        ])
        result = extract_contract_from_docx(docx_bytes, "generic_doc.docx", mode="mock")

        assert result.is_valid is True
        # Should have defaults
        assert result.contract_data["contract_type"] == "FFP"  # default
        assert result.report.defaulted > 0
        assert len(result.report.warnings) > 0


# ══════════════════════════════════════════════════════════════════════════════
# UPLOAD MANAGER TESTS
# ══════════════════════════════════════════════════════════════════════════════

class TestUploadManager:
    """Tests for .docx support in upload_manager."""

    def test_validate_upload_txt(self):
        from src.ingestion.upload_manager import validate_upload

        # Valid .txt
        valid_txt = (
            "CONTRACT NUMBER: TEST-001\n"
            "--- APPLICABLE FAR CLAUSES ---\n"
            "FAR 52.204-21\n" * 10
        )
        is_valid, _ = validate_upload(valid_txt, "test.txt")
        assert is_valid is True

    def test_validate_upload_txt_invalid(self):
        from src.ingestion.upload_manager import validate_upload

        is_valid, error = validate_upload("short", "test.txt")
        assert is_valid is False

    def test_validate_upload_docx(self):
        from src.ingestion.upload_manager import validate_upload

        docx_bytes = _create_contract_docx()
        is_valid, error = validate_upload(docx_bytes, "contract.docx")
        assert is_valid is True

    def test_validate_upload_docx_invalid(self):
        from src.ingestion.upload_manager import validate_upload

        is_valid, error = validate_upload(b"not a docx", "bad.docx")
        assert is_valid is False

    def test_save_uploaded_file_bytes(self):
        from src.ingestion.upload_manager import save_uploaded_file

        docx_bytes = _create_test_docx(paragraphs=["test"])
        with tempfile.TemporaryDirectory() as tmpdir:
            path = save_uploaded_file(docx_bytes, "test.docx", tmpdir)
            assert path.exists()
            assert path.read_bytes() == docx_bytes

    def test_save_uploaded_file_dedup(self):
        from src.ingestion.upload_manager import save_uploaded_file

        with tempfile.TemporaryDirectory() as tmpdir:
            # Save first file
            save_uploaded_file("content1", "test.txt", tmpdir)
            # Save second with same name — should get timestamped
            path2 = save_uploaded_file("content2", "test.txt", tmpdir)
            assert "test_" in path2.name  # Should have timestamp


# ══════════════════════════════════════════════════════════════════════════════
# INTEGRATION TEST
# ══════════════════════════════════════════════════════════════════════════════

class TestIntegration:
    """End-to-end test: .docx → extraction → chunking → LangChain documents."""

    def test_docx_to_langchain_documents(self):
        """Full pipeline: create docx → extract → chunk → LangChain docs."""
        from src.ingestion.hybrid_extractor import extract_contract_from_docx
        from src.ingestion.chunker import chunk_contract, chunks_to_langchain_documents

        docx_bytes = _create_contract_docx()
        result = extract_contract_from_docx(docx_bytes, "integration_test.docx", mode="mock")

        assert result.is_valid
        assert result.document_text

        # Chunk the generated text
        chunks = chunk_contract(result.document_text)
        assert len(chunks) > 0

        # Convert to LangChain Documents
        documents = chunks_to_langchain_documents(chunks)
        assert len(documents) > 0

        # Verify metadata propagation
        for doc in documents:
            assert doc.metadata.get("contract_number") == "W911NF-25-D-0042"
            assert doc.page_content  # Non-empty content

    def test_extraction_report_serializable(self):
        """Ensure the extraction report can be serialized to JSON."""
        from src.ingestion.upload_manager import _report_to_dict
        from src.ingestion.hybrid_extractor import extract_contract_from_docx

        docx_bytes = _create_contract_docx()
        result = extract_contract_from_docx(docx_bytes, "test.docx", mode="mock")

        report_dict = _report_to_dict(result.report)
        # Should be JSON-serializable
        json_str = json.dumps(report_dict)
        assert json_str
        parsed = json.loads(json_str)
        assert "total_fields" in parsed
        assert "confidence_avg" in parsed
